#!/usr/bin/env python3
"""
Word Document Generator for Piano Keys Production Readiness Plan
Generates a professionally formatted .docx document with styling, tables, and checkboxes.
"""

from pathlib import Path
from datetime import date, datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def create_word_document(source_file: Path, output_dir: Path) -> Path:
    """
    Generate a professionally formatted Word document from the production readiness plan.

    Args:
        source_file: Path to the source markdown file
        output_dir: Directory to save the generated document

    Returns:
        Path to the generated .docx file
    """
    # Create document
    doc = Document()

    # Configure document properties
    doc.core_properties.title = "Piano Keys - Production Readiness Plan"
    doc.core_properties.subject = "12-Week Roadmap to Production Launch"
    doc.core_properties.author = "Piano Keys Team"
    doc.core_properties.created = datetime.now()

    # Read source content
    with open(source_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Apply professional styling
    apply_document_styles(doc)

    # Add cover page
    add_cover_page(doc)
    doc.add_page_break()

    # Parse and add content
    add_content_from_markdown(doc, content)

    # Save document
    output_path = output_dir / "production_readiness_plan.docx"
    doc.save(output_path)

    return output_path


def apply_document_styles(doc):
    """Apply professional styling to the document."""
    # Title style
    styles = doc.styles

    # Heading 1 - Blue, large
    heading1 = styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(24)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(31, 71, 136)  # #1F4788

    # Heading 2 - Darker blue, medium
    heading2 = styles['Heading 2']
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(18)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(46, 92, 138)  # #2E5C8A

    # Heading 3 - Bold, black
    heading3 = styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(14)
    heading3.font.bold = True

    # Normal text
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)


def add_cover_page(doc):
    """Add a professional cover page."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Piano Keys")
    title_run.font.size = Pt(44)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 71, 136)

    # Subtitle
    doc.add_paragraph()  # Spacing
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Production Readiness Plan")
    subtitle_run.font.size = Pt(28)
    subtitle_run.font.color.rgb = RGBColor(46, 92, 138)

    # Subheading
    doc.add_paragraph()  # Spacing
    subheading = doc.add_paragraph()
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading_run = subheading.add_run("12-Week Roadmap to Launch")
    subheading_run.font.size = Pt(18)
    subheading_run.font.italic = True

    # Add spacing
    for _ in range(3):
        doc.add_paragraph()

    # Date and version
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Document Date: {date.today().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)

    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_run = version_para.add_run("Version 1.0")
    version_run.font.size = Pt(12)


def add_content_from_markdown(doc, content):
    """Parse markdown and add formatted content to the document."""
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Handle headers
        if line.startswith('# '):
            # H1 - skip if it's the title (already on cover)
            if i > 5:  # Not in header area
                heading_text = line.lstrip('#').strip()
                p = doc.add_heading(heading_text, level=1)
                p.style = 'Heading 1'
            i += 1

        elif line.startswith('## '):
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=2)
            i += 1

        elif line.startswith('### '):
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=3)
            i += 1

        # Handle horizontal rules
        elif line.startswith('---'):
            doc.add_paragraph()  # Add spacing
            i += 1

        # Handle list items
        elif line.startswith('- '):
            text = line.lstrip('-').strip()
            # Check if it's a checkbox
            if text.startswith('[ ]'):
                text = '☐ ' + text[3:].strip()
            elif text.startswith('[x]') or text.startswith('[X]'):
                text = '☑ ' + text[3:].strip()

            p = doc.add_paragraph(text, style='List Bullet')
            i += 1

        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            i += 1

        # Handle code blocks
        elif line.startswith('```'):
            # Start of code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1

            # Add code block
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                # Style as code
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    # Add gray background would require more complex styling

            i += 1  # Skip closing ```

        # Handle bold text
        elif line.startswith('**') and line.endswith('**'):
            text = line.strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            i += 1

        # Regular paragraph
        else:
            # Check if next lines are part of same paragraph
            paragraph_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not (
                lines[i].startswith('#') or
                lines[i].startswith('-') or
                lines[i].startswith('```') or
                lines[i].startswith('---') or
                re.match(r'^\d+\.', lines[i])
            ):
                paragraph_lines.append(lines[i].strip())
                i += 1

            full_text = ' '.join(paragraph_lines)

            # Add paragraph with inline formatting
            p = doc.add_paragraph()
            add_formatted_text(p, full_text)

    # Add footer
    add_footer(doc)


def add_formatted_text(paragraph, text):
    """Add text to paragraph with inline formatting (bold, code, etc.)."""
    # Simple implementation - add full text
    # For production, you'd want to parse **bold**, `code`, etc.
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))', text)

    for part in parts:
        if not part:
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part.strip('*'))
            run.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Code
            run = paragraph.add_run(part.strip('`'))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # Normal text
            paragraph.add_run(part)


def add_footer(doc):
    """Add footer to the document."""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Piano Keys Production Readiness Plan"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0]
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)


if __name__ == "__main__":
    # Test generation
    source = Path(__file__).parent.parent / ".claude/plans/production-readiness-plan-source.md"
    output = Path(__file__).parent.parent / "output"
    output.mkdir(exist_ok=True)

    result = create_word_document(source, output)
    print(f"✅ Word document generated: {result}")
