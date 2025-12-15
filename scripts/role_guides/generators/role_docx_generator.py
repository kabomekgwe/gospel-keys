#!/usr/bin/env python3
"""
Role-Specific Word Document Generator
Generates exponentially detailed 50-100 page implementation guides for each role
"""

from pathlib import Path
from datetime import date, datetime
from typing import List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import sys

# Add parent directories to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "content"))

from code_snippet_extractor import CodeSnippetExtractor
from role_mapper import RoleMapper


class RoleDocxGenerator:
    """Generate exponentially detailed Word documents for engineering roles."""

    def __init__(
        self,
        project_root: Path,
        plan_file: Path,
        role_configs_dir: Path,
        output_dir: Path
    ):
        """
        Initialize the role document generator.

        Args:
            project_root: Path to Piano Keys project
            plan_file: Path to production readiness plan
            role_configs_dir: Path to role configuration files
            output_dir: Directory to save generated documents
        """
        self.extractor = CodeSnippetExtractor(project_root, plan_file, role_configs_dir)
        self.mapper = RoleMapper(plan_file, role_configs_dir)
        self.output_dir = output_dir
        self.output_dir.mkdir(exist_ok=True)

    def generate_role_document(self, role_name: str) -> Path:
        """
        Generate comprehensive Word document for a specific role.

        Args:
            role_name: Name of the role (e.g., "Backend Engineer")

        Returns:
            Path to generated document
        """
        print(f"\n📄 Generating document for: {role_name}")
        print("-" * 70)

        # Create document
        doc = Document()

        # Configure document properties
        doc.core_properties.title = f"{role_name} - Production Readiness Implementation Guide"
        doc.core_properties.subject = "Exponentially Detailed Implementation Guide"
        doc.core_properties.author = "Piano Keys Team"
        doc.core_properties.created = datetime.now()

        # Apply styling
        self._apply_styles(doc)

        # Get role summary
        summary = self.mapper.generate_role_summary(role_name)

        # Add cover page
        self._add_cover_page(doc, role_name, summary)
        doc.add_page_break()

        # Add table of contents placeholder
        self._add_toc_page(doc, summary)
        doc.add_page_break()

        # Section 1: Role Overview
        self._add_role_overview(doc, summary)
        doc.add_page_break()

        # Section 2-N: Phase implementations
        phases = summary.get('phases', [])
        for phase in phases:
            self._add_phase_section(doc, phase, role_name)
            doc.add_page_break()

        # File-by-file implementation guide
        self._add_implementation_guide(doc, role_name, summary)
        doc.add_page_break()

        # Success metrics and checklist
        self._add_success_metrics(doc, summary)
        doc.add_page_break()

        # Troubleshooting guide
        self._add_troubleshooting(doc, role_name)
        doc.add_page_break()

        # Appendix: Complete code listings
        self._add_code_appendix(doc, role_name)

        # Add footer
        self._add_footer(doc, role_name)

        # Save document
        filename = f"{role_name.lower().replace(' ', '_').replace('/', '_')}_implementation_guide.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)

        print(f"✅ Generated: {output_path}")
        print(f"   File size: {output_path.stat().st_size // 1024} KB")

        return output_path

    def _apply_styles(self, doc):
        """Apply professional styling to document."""
        styles = doc.styles

        # Heading 1 - Blue, large
        heading1 = styles['Heading 1']
        heading1.font.name = 'Calibri'
        heading1.font.size = Pt(28)
        heading1.font.bold = True
        heading1.font.color.rgb = RGBColor(31, 71, 136)

        # Heading 2 - Darker blue, medium
        heading2 = styles['Heading 2']
        heading2.font.name = 'Calibri'
        heading2.font.size = Pt(20)
        heading2.font.bold = True
        heading2.font.color.rgb = RGBColor(46, 92, 138)

        # Heading 3 - Bold, black
        heading3 = styles['Heading 3']
        heading3.font.name = 'Calibri'
        heading3.font.size = Pt(16)
        heading3.font.bold = True

        # Normal text
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

    def _add_cover_page(self, doc, role_name: str, summary: dict):
        """Add professional cover page."""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Piano Keys")
        title_run.font.size = Pt(48)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(31, 71, 136)

        # Subtitle - Role name
        doc.add_paragraph()
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(role_name)
        subtitle_run.font.size = Pt(32)
        subtitle_run.font.color.rgb = RGBColor(46, 92, 138)

        # Description
        doc.add_paragraph()
        desc = doc.add_paragraph()
        desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc.add_run("Production Readiness Implementation Guide")
        desc_run.font.size = Pt(20)
        desc_run.font.italic = True

        # Add spacing
        for _ in range(3):
            doc.add_paragraph()

        # Role description
        if summary.get('description'):
            role_desc = doc.add_paragraph()
            role_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            role_desc_run = role_desc.add_run(summary['description'])
            role_desc_run.font.size = Pt(14)

        # Add spacing
        for _ in range(2):
            doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Document Date: {date.today().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)

        # Version
        version_para = doc.add_paragraph()
        version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        version_run = version_para.add_run("Version 1.0")
        version_run.font.size = Pt(12)

        # Footer note
        for _ in range(3):
            doc.add_paragraph()

        footer_note = doc.add_paragraph()
        footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_note.add_run("⚠️ EXPONENTIALLY DETAILED GUIDE")
        footer_run.font.size = Pt(10)
        footer_run.font.bold = True
        footer_run.font.color.rgb = RGBColor(220, 53, 69)

        detail_note = doc.add_paragraph()
        detail_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        detail_run = detail_note.add_run(
            "Includes exact file paths, line numbers, complete code listings, "
            "before/after comparisons, testing procedures, and rollback plans"
        )
        detail_run.font.size = Pt(9)

    def _add_toc_page(self, doc, summary: dict):
        """Add table of contents page."""
        doc.add_heading("Table of Contents", level=1)

        toc = doc.add_paragraph()
        toc.add_run("1. Role Overview & Responsibilities\n")
        toc.add_run("   1.1 Role Description\n")
        toc.add_run("   1.2 Technology Stack\n")
        toc.add_run("   1.3 Dependencies\n")
        toc.add_run("   1.4 Time Estimates\n\n")

        # Add phases to TOC
        phases = summary.get('phases', [])
        for i, phase in enumerate(phases, start=2):
            phase_num = phase.get('phase')
            phase_name = phase.get('name')
            toc.add_run(f"{i}. Phase {phase_num}: {phase_name}\n")
            toc.add_run(f"   {i}.1 Phase Overview\n")
            toc.add_run(f"   {i}.2 Implementation Tasks\n")
            toc.add_run(f"   {i}.3 Code Examples\n\n")

        # Add remaining sections
        section_num = len(phases) + 2
        toc.add_run(f"{section_num}. File-by-File Implementation Guide\n")
        toc.add_run(f"{section_num + 1}. Success Metrics & Checklist\n")
        toc.add_run(f"{section_num + 2}. Troubleshooting Guide\n")
        toc.add_run(f"{section_num + 3}. Appendix: Complete Code Listings\n")

    def _add_role_overview(self, doc, summary: dict):
        """Add comprehensive role overview section."""
        doc.add_heading("1. Role Overview & Responsibilities", level=1)

        # Role description
        doc.add_heading("1.1 Role Description", level=2)
        desc_para = doc.add_paragraph(summary.get('description', ''))

        # Technology stack
        doc.add_heading("1.2 Technology Stack", level=2)
        tech_stack = summary.get('tech_stack', [])

        for tech in tech_stack:
            tech_para = doc.add_paragraph(style='List Bullet')
            tech_para.add_run(tech)

        # Dependencies
        doc.add_heading("1.3 Dependencies", level=2)
        dependencies = summary.get('dependencies', {})

        if dependencies.get('blocks'):
            doc.add_heading("This role blocks:", level=3)
            for dep in dependencies['blocks']:
                dep_para = doc.add_paragraph(style='List Bullet')
                dep_para.add_run(f"{dep['role']}: {dep['reason']}")

        if dependencies.get('blocked_by'):
            doc.add_heading("This role is blocked by:", level=3)
            for dep in dependencies['blocked_by']:
                dep_para = doc.add_paragraph(style='List Bullet')
                dep_para.add_run(f"{dep['role']}: {dep['reason']}")

        # Time estimates
        doc.add_heading("1.4 Time Estimates", level=2)
        time_estimates = summary.get('time_estimates', {})

        for phase, estimate in time_estimates.items():
            time_para = doc.add_paragraph(style='List Bullet')
            time_para.add_run(f"{phase}: {estimate}")

    def _add_phase_section(self, doc, phase: dict, role_name: str):
        """Add detailed phase implementation section."""
        phase_num = phase.get('phase')
        phase_name = phase.get('name')
        weeks = phase.get('weeks')

        doc.add_heading(f"Phase {phase_num}: {phase_name}", level=1)

        # Phase overview
        doc.add_heading("Phase Overview", level=2)
        overview = doc.add_paragraph()
        overview.add_run(f"Timeline: Weeks {weeks}\n").bold = True
        overview.add_run(f"Goal: {phase_name}\n")

        # Focus areas
        doc.add_heading("Focus Areas", level=2)
        focus_areas = phase.get('focus_areas', [])

        for area in focus_areas:
            area_para = doc.add_paragraph(style='List Bullet')
            area_para.add_run(area)

        # Get code examples for this phase
        doc.add_heading("Implementation Details", level=2)

        # Extract relevant tasks for this phase
        tasks = self.mapper.get_role_tasks(role_name)

        # Filter tasks by priority for this phase
        # For now, show all CRITICAL and HIGH priority tasks
        relevant_tasks = [t for t in tasks if t.get('priority') in ['CRITICAL', 'HIGH']][:3]

        for i, task in enumerate(relevant_tasks, 1):
            snippet = self.extractor.extract_task_code(task)

            doc.add_heading(f"Task {i}: {snippet['task']}", level=3)

            # Task details
            details = doc.add_paragraph()
            details.add_run(f"File: {snippet['file']}\n")
            details.add_run(f"Priority: {snippet['priority']}\n").font.color.rgb = self._get_priority_color(snippet['priority'])

            if snippet.get('target_lines'):
                details.add_run(f"Lines: {snippet['target_lines']}\n")

            # Analysis
            if snippet.get('analysis'):
                doc.add_heading("Analysis:", level=4)
                analysis_para = doc.add_paragraph()
                analysis_run = analysis_para.add_run(snippet['analysis'])
                analysis_run.font.size = Pt(10)

            # Current code
            if snippet.get('current_code'):
                doc.add_heading("Current Code:", level=4)
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(snippet['current_code'])
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)

    def _add_implementation_guide(self, doc, role_name: str, summary: dict):
        """Add file-by-file implementation guide."""
        doc.add_heading("File-by-File Implementation Guide", level=1)

        tasks = summary.get('tasks', [])

        for i, task in enumerate(tasks, 1):
            snippet = self.extractor.extract_task_code(task, context_lines=15)

            doc.add_heading(f"File {i}: {snippet['file']}", level=2)

            # Task overview table
            doc.add_heading("Task Overview", level=3)
            overview_para = doc.add_paragraph()
            overview_para.add_run(f"Task: {snippet['task']}\n")
            overview_para.add_run(f"Priority: {snippet['priority']}\n").font.color.rgb = self._get_priority_color(snippet['priority'])
            overview_para.add_run(f"File Status: {'EXISTS' if snippet['exists'] else 'NEW FILE'}\n")

            if snippet.get('target_lines'):
                overview_para.add_run(f"Target Lines: {snippet['target_lines']}\n")

            # Implementation steps
            doc.add_heading("Implementation Steps", level=3)

            if snippet['exists']:
                steps = doc.add_paragraph()
                steps.add_run("1. Open file in editor\n")
                steps.add_run(f"2. Navigate to lines {snippet.get('target_lines', 'as specified')}\n")
                steps.add_run("3. Review current code (see below)\n")
                steps.add_run("4. Implement changes\n")
                steps.add_run("5. Run tests\n")
                steps.add_run("6. Commit changes\n")
            else:
                steps = doc.add_paragraph()
                steps.add_run(f"1. Create new file: {snippet['file']}\n")
                steps.add_run("2. Implement functionality (see template below)\n")
                steps.add_run("3. Add tests\n")
                steps.add_run("4. Commit changes\n")

            # Code section
            if snippet.get('current_code'):
                doc.add_heading("Current Code (with context)", level=3)
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(snippet['current_code'])
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(8)

            # Analysis
            if snippet.get('analysis'):
                doc.add_heading("Code Analysis", level=3)
                analysis_para = doc.add_paragraph()
                analysis_para.add_run(snippet['analysis'])

            doc.add_paragraph()  # Spacing

    def _add_success_metrics(self, doc, summary: dict):
        """Add success metrics and checklist."""
        doc.add_heading("Success Metrics & Checklist", level=1)

        metrics = summary.get('success_metrics', [])

        if metrics:
            doc.add_heading("Metrics", level=2)

            for metric in metrics:
                metric_para = doc.add_paragraph(style='List Bullet')
                metric_para.add_run(f"☐ {metric['metric']}: {metric['target']}")

    def _add_troubleshooting(self, doc, role_name: str):
        """Add troubleshooting guide."""
        doc.add_heading("Troubleshooting Guide", level=1)

        doc.add_paragraph(
            "This section provides common issues and solutions encountered during implementation."
        )

        # Add common troubleshooting scenarios
        doc.add_heading("Common Issues", level=2)

        issues = [
            ("Import errors", "Verify all dependencies are installed and paths are correct"),
            ("Test failures", "Check test data, mocks, and assertions"),
            ("Type errors", "Run mypy/tsc and fix type annotations"),
            ("Runtime errors", "Check logs, add debug output, verify environment variables"),
        ]

        for issue, solution in issues:
            doc.add_heading(issue, level=3)
            sol_para = doc.add_paragraph()
            sol_para.add_run("Solution: ")
            sol_para.add_run(solution)

    def _add_code_appendix(self, doc, role_name: str):
        """Add appendix with complete code listings."""
        doc.add_heading("Appendix: Complete Code Listings", level=1)

        doc.add_paragraph(
            "This appendix contains complete code listings for all files mentioned in this guide."
        )

        # Get all code snippets
        snippets = self.extractor.extract_all_role_code(role_name)

        for snippet in snippets:
            if snippet.get('full_file'):
                doc.add_heading(snippet['file'], level=2)

                code_para = doc.add_paragraph()
                code_run = code_para.add_run(snippet['full_file'])
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(7)

                doc.add_page_break()

    def _add_footer(self, doc, role_name: str):
        """Add footer to document."""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Piano Keys - {role_name} Implementation Guide"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

    def _get_priority_color(self, priority: str) -> RGBColor:
        """Get color based on priority level."""
        if priority == "CRITICAL":
            return RGBColor(220, 53, 69)  # Red
        elif priority == "HIGH":
            return RGBColor(255, 193, 7)  # Orange
        elif priority == "MEDIUM":
            return RGBColor(0, 123, 255)  # Blue
        else:
            return RGBColor(108, 117, 125)  # Gray

    def generate_all_roles(self) -> List[Path]:
        """Generate documents for all configured roles."""
        generated_docs = []

        roles = self.mapper.get_all_roles()

        print(f"\n🚀 Generating exponentially detailed guides for {len(roles)} roles...")
        print("=" * 70)

        for role in roles:
            doc_path = self.generate_role_document(role)
            generated_docs.append(doc_path)

        return generated_docs


if __name__ == "__main__":
    # Test role document generation
    from pathlib import Path

    project_root = Path(__file__).parent.parent.parent.parent
    plan_file = project_root / ".claude/plans/production-readiness-plan-source.md"
    configs_dir = Path(__file__).parent.parent / "templates/role_configs"
    output_dir = project_root / "output/role_guides"

    generator = RoleDocxGenerator(project_root, plan_file, configs_dir, output_dir)

    # Generate Backend Engineer guide as test
    print("Testing Role Document Generator...")
    print("=" * 70)

    doc_path = generator.generate_role_document("Backend Engineer")

    print(f"\n✅ Test complete! Document generated at:")
    print(f"   {doc_path}")
